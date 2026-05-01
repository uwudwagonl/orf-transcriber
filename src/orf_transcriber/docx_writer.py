"""Produce a structured Word document from a transcription result.

Layout (German UI labels, AI-friendly section markers):

  Title (Heading 1)
  Metadata table

  ## Inhaltsverzeichnis (Heading 2)
  Plain prose, no timestamps. Best for AI summarisation.

  ## Transkript mit Zeitmarken (Heading 2)
  Each segment as `[hh:mm:ss] text` paragraph.

  ## Maschinenlesbare Zusammenfassung (Heading 2)
  Hidden-style key/value list to help LLMs extract structure.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt

from .metadata import VideoMeta
from .transcriber import TranscriptionResult


def _fmt_ts(seconds: float) -> str:
    s = int(round(seconds))
    h, rem = divmod(s, 3600)
    m, s = divmod(rem, 60)
    return f"{h:02d}:{m:02d}:{s:02d}"


def _fmt_duration(seconds: float) -> str:
    if seconds <= 0:
        return "—"
    s = int(round(seconds))
    h, rem = divmod(s, 3600)
    m, s = divmod(rem, 60)
    if h:
        return f"{h} h {m:02d} min"
    if m:
        return f"{m} min {s:02d} s"
    return f"{s} s"


def write_transcript(
    output_path: Path,
    meta: VideoMeta,
    result: TranscriptionResult,
    source_url: str,
    video_file: Path | None,
) -> Path:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(meta.title or "ORF-Video Transkript", level=1)

    # ── metadata table ────────────────────────────────────────────────────
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Light List"

    rows: list[tuple[str, str]] = [
        ("Quelle", source_url),
        ("Veröffentlicht", meta.published or "—"),
        ("Sprache (erkannt)", result.language.upper()),
        ("Dauer", _fmt_duration(result.duration)),
        ("Transkribiert am", datetime.now().strftime("%d.%m.%Y %H:%M")),
    ]
    if video_file:
        rows.append(("Videodatei", str(video_file)))

    for key, val in rows:
        row = table.add_row().cells
        row[0].text = key
        row[1].text = val
        row[0].paragraphs[0].runs[0].bold = True

    if meta.description:
        doc.add_paragraph()
        doc.add_paragraph(meta.description, style="Intense Quote")

    doc.add_paragraph()

    # ── plain prose section ──────────────────────────────────────────────
    doc.add_heading("Transkript (Fließtext)", level=2)
    doc.add_paragraph(
        "Diese Fassung enthält keine Zeitmarken und eignet sich am besten zum "
        "Vorlesen, Lesen oder zum Einlesen in eine KI für Zusammenfassungen.",
        style="Intense Quote",
    )

    full_text = " ".join(seg.text for seg in result.segments).strip()
    if full_text:
        # Break into ~5-segment paragraphs for readability.
        chunk: list[str] = []
        for i, seg in enumerate(result.segments, start=1):
            chunk.append(seg.text)
            if i % 5 == 0:
                doc.add_paragraph(" ".join(chunk))
                chunk = []
        if chunk:
            doc.add_paragraph(" ".join(chunk))
    else:
        doc.add_paragraph("(Kein Text erkannt.)")

    # ── timestamped section ──────────────────────────────────────────────
    doc.add_heading("Transkript mit Zeitmarken", level=2)
    doc.add_paragraph(
        "Jeder Absatz beginnt mit dem Zeitpunkt im Video, an dem der Satz "
        "gesprochen wird (Stunden:Minuten:Sekunden).",
        style="Intense Quote",
    )

    for seg in result.segments:
        para = doc.add_paragraph()
        run_ts = para.add_run(f"[{_fmt_ts(seg.start)}] ")
        run_ts.bold = True
        para.add_run(seg.text)

    # ── machine-readable footer ──────────────────────────────────────────
    doc.add_heading("Maschinenlesbare Zusammenfassung", level=2)
    doc.add_paragraph(
        "Dieser Abschnitt ist für KI-Werkzeuge gedacht und enthält "
        "strukturierte Metadaten in einem einfachen Schlüssel/Wert-Format.",
        style="Intense Quote",
    )

    machine_lines = [
        "BEGIN ORF_TRANSCRIPT_METADATA",
        f"title: {meta.title}",
        f"source_url: {source_url}",
        f"canonical_url: {meta.canonical_url}",
        f"published: {meta.published}",
        f"language_detected: {result.language}",
        f"duration_seconds: {result.duration:.2f}",
        f"segment_count: {len(result.segments)}",
        f"transcribed_at: {datetime.now().isoformat(timespec='seconds')}",
        "END ORF_TRANSCRIPT_METADATA",
    ]
    for line in machine_lines:
        p = doc.add_paragraph(line)
        p.style = doc.styles["No Spacing"] if "No Spacing" in [s.name for s in doc.styles] else p.style

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
